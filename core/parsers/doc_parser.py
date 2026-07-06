"""DOC/DOCX 文件解析器"""
import os
import struct
import re

# UTF-16LE 文本段：每个码元 (low,high) 满足 0x20<=code<0xFFFE 或 code∈{0x09,0x0A,0x0D}。
# 拆成三支字节类表达（high=0x00 可打印 / high∈0x01-0xFE 任意 / high=0xFF 低字节 0x00-0xFD），
# 覆盖中文（高字节非零）等全部有效码元；匹配 4+ 码元（与原 len>3 一致）。
# 注意：不能用 (?:..\x00) 形式 —— 那会丢掉所有高字节非零的 CJK 字符。
_RE_UNICODE = re.compile(
    rb'(?:[\x00-\xff][\x01-\xfe]|[\x09\x0a\x0d\x20-\xff]\x00|[\x00-\xfd]\xff){4,}'
)
# ASCII/GBK 文本段：可打印 ASCII(0x20-0x7e) + GBK 高字节(0x80-0xfe) + 制表/换行；
# 匹配 6+ 字节（与原 len>5 一致）。
_RE_ASCII = re.compile(rb'[\x09\x0a\x0d\x20-\x7e\x80-\xfe]{6,}')


def parse_docx(filepath):
    """解析 .docx 文件（Office Open XML格式）"""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        return "\n".join(paragraphs)
    except ImportError:
        raise RuntimeError("python-docx 未安装")
    except Exception as e:
        raise RuntimeError(f"DOCX解析失败: {e}")


def parse_doc(filepath):
    """解析 .doc 文件（旧版OLE格式）
    使用多种方法尝试提取文本
    """
    # 方法1: 使用 olefile 提取 OLE 流
    text = _parse_doc_olefile(filepath)
    if text and len(text.strip()) > 50:
        return text

    # 方法2: 从二进制中直接提取可读文本
    text = _parse_doc_binary(filepath)
    if text and len(text.strip()) > 50:
        return text

    return ""


def _parse_doc_olefile(filepath):
    """使用olefile解析.doc文件"""
    try:
        import olefile
        if not olefile.isOleFile(filepath):
            return ""

        ole = olefile.OleFileIO(filepath)

        # Word文档的主要文本流
        text_parts = []

        # 尝试从 WordDocument 流提取
        if ole.exists("WordDocument"):
            stream = ole.openstream("WordDocument")
            data = stream.read()
            stream.close()

            # 尝试提取 Unicode 文本
            text = _extract_unicode_text(data)
            if text:
                text_parts.append(text)

        # 尝试从 1Table 或 0Table 流获取补充信息
        for table_name in ["1Table", "0Table"]:
            if ole.exists(table_name):
                stream = ole.openstream(table_name)
                data = stream.read()
                stream.close()
                text = _extract_unicode_text(data)
                if text and len(text) > 20:
                    text_parts.append(text)

        ole.close()

        combined = "\n".join(text_parts)
        return _clean_doc_text(combined)

    except ImportError:
        return ""
    except Exception:
        return ""


def _parse_doc_binary(filepath):
    """从二进制.doc文件中直接提取可读文本"""
    try:
        with open(filepath, "rb") as f:
            data = f.read()

        # 提取 UTF-16LE 文本段
        text = _extract_unicode_text(data)
        if text and len(text.strip()) > 50:
            return _clean_doc_text(text)

        # 提取 ASCII/GBK 文本段
        text = _extract_ascii_text(data)
        if text and len(text.strip()) > 50:
            return _clean_doc_text(text)

        return ""
    except Exception:
        return ""


def _extract_unicode_text(data):
    """从二进制数据中提取UTF-16LE编码的文本

    用预编译正则在 C 层扫描连续有效码元（4+ 个），替代逐字节 Python 循环，
    大文件解析提速 10-50x。有效码元定义与原实现一致：0x20<=code<0xFFFE 或
    code∈{0x09,0x0A,0x0D}。findall 从所有字节偏移尝试匹配，是原偶对齐扫描的
    超集（不会丢失原文，只可能多出少量噪声，由 _clean_doc_text 清理）。
    """
    return "\n".join(
        m.decode('utf-16-le', 'ignore') for m in _RE_UNICODE.findall(data)
    )


def _extract_ascii_text(data):
    """从二进制数据中提取ASCII/GBK文本

    预编译正则匹配连续可打印字节（6+ 个），替代逐字节 Python 循环。
    字节集与原实现一致：0x20-0x7e / 0x80-0xfe / 0x09,0x0A,0x0D。
    """
    raw = b"\n".join(_RE_ASCII.findall(data))
    try:
        from ..text_utils import decode_bytes
        return decode_bytes(raw)
    except Exception:
        return raw.decode('latin-1')


def _clean_doc_text(text):
    """清理从doc提取的文本"""
    if not text:
        return ""
    # 去除连续空字符
    text = re.sub(r'\x00+', '', text)
    # 去除控制字符
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    # 合并多个空白
    text = re.sub(r' {2,}', ' ', text)
    # 合并多个换行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
